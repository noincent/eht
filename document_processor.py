from docx import Document
import re
import hashlib
import json
from typing import Dict, List, Tuple, Optional, Any
import numpy as np
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModel
import torch
import torch.nn.functional as F

class LanguageDetector:
    """Detects document language and processes single-language documents."""
    
    @staticmethod
    def detect_language(text: str) -> str:
        """Detect if text is primarily Chinese or English."""
        # Simple heuristic: count Chinese characters vs ASCII
        chinese_char_count = len(re.findall(r'[\u4e00-\u9fff]', text))
        ascii_char_count = len(re.findall(r'[a-zA-Z]', text))
        
        return 'zh' if chinese_char_count > ascii_char_count else 'en'
    
    @staticmethod
    def process_document(doc_path: str, language: str) -> Dict:
        """Process a single-language document and extract sections."""
        doc = Document(doc_path)
        sections = {}
        
        current_id = ""
        current_content = []
        
        # Pattern based on language
        if language == 'zh':
            header_pattern = r'^(第[一二三四五六七八九十]+章|[一二三四五六七八九十]+、|\d+[\.、]|[（(]\d+[）)])'
        else:  # English
            header_pattern = r'^(Chapter|Section|Article|Part)\s+[\w\d\.]+|^\d+[\.\)]\s+[A-Z]'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a section header
            is_header = bool(re.match(header_pattern, text))
            
            if is_header:
                # Save previous section if it exists
                if current_id and current_content:
                    sections[current_id] = {
                        'title': current_id,
                        'content': ' '.join(current_content)
                    }
                
                # Start new section
                current_id = text
                current_content = []
            else:
                # Add to current section
                current_content.append(text)
                
        # Don't forget the last section
        if current_id and current_content:
            sections[current_id] = {
                'title': current_id,
                'content': ' '.join(current_content)
            }
                
        return sections


class RecursiveChunker:
    """Chunks documents at multiple granularities with sliding window."""
    
    def __init__(self, 
                 max_chunk_size: int = 512,
                 min_chunk_size: int = 50,
                 chunk_overlap: int = 50):
        self.max_chunk_size = max_chunk_size
        self.min_chunk_size = min_chunk_size
        self.chunk_overlap = chunk_overlap
    
    def create_chunks(self, document: Dict, language: str) -> List[Dict]:
        """
        Create multiple levels of chunks from the document.
        Returns a list of chunks with metadata.
        """
        chunks = []
        
        # First level: entire sections
        for section_id, section in document.items():
            section_text = f"{section['title']} {section['content']}"
            section_hash = self._get_hash(section_text)
            
            chunks.append({
                'id': section_hash,
                'text': section_text,
                'metadata': {
                    'level': 'section',
                    'section_id': section_id,
                    'language': language,
                    'parent_id': None
                }
            })
            
            # Second level: paragraphs
            paragraphs = self._split_into_paragraphs(section['content'])
            for i, para in enumerate(paragraphs):
                if len(para) < self.min_chunk_size:
                    continue
                    
                para_hash = self._get_hash(para)
                chunks.append({
                    'id': para_hash,
                    'text': para,
                    'metadata': {
                        'level': 'paragraph',
                        'section_id': section_id,
                        'paragraph_idx': i,
                        'language': language,
                        'parent_id': section_hash
                    }
                })
            
            # Third level: sliding window for long sections
            if len(section['content']) > self.max_chunk_size:
                sliding_chunks = self._create_sliding_chunks(section['content'])
                for i, chunk_text in enumerate(sliding_chunks):
                    chunk_hash = self._get_hash(chunk_text)
                    chunks.append({
                        'id': chunk_hash,
                        'text': chunk_text,
                        'metadata': {
                            'level': 'sliding_window',
                            'section_id': section_id,
                            'chunk_idx': i,
                            'language': language,
                            'parent_id': section_hash
                        }
                    })
        
        return chunks
    
    def _get_hash(self, text: str) -> str:
        """Create a unique hash for a text chunk."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Split text into paragraphs."""
        return [p for p in re.split(r'\n+', text) if p.strip()]
    
    def _create_sliding_chunks(self, text: str) -> List[str]:
        """Create overlapping chunks using sliding window."""
        tokens = text.split()
        chunks = []
        
        if len(tokens) <= self.max_chunk_size:
            return [text]
        
        for i in range(0, len(tokens), self.max_chunk_size - self.chunk_overlap):
            chunk = ' '.join(tokens[i:i + self.max_chunk_size])
            if len(chunk.split()) >= self.min_chunk_size:
                chunks.append(chunk)
                
        return chunks


class EnglishEmbedder:
    """Handles embedding generation for English text."""
    
    def __init__(self, model_name: str = "intfloat/e5-large-v2"):
        self.model = SentenceTransformer(model_name)
        
    def get_embeddings(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings for a list of texts."""
        return self.model.encode(texts, show_progress_bar=True)


class ChineseEmbedder:
    """Handles embedding generation for Chinese text."""
    
    def __init__(self, model_name: str = "shibing624/text2vec-base-chinese"):
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer(model_name)
            self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
            self.model.to(self.device)
            print(f"✅ Loaded Chinese embedding model: {model_name}")
        except Exception as e:
            print(f"Error loading Chinese embedding model: {e}")
            # Fallback to English model if Chinese model fails
            fallback_model = "intfloat/e5-base"
            print(f"⚠️ Falling back to multilingual model: {fallback_model}")
            self.model = SentenceTransformer(fallback_model)
            self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
            self.model.to(self.device)
        
    def get_embeddings(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings for a list of Chinese texts."""
        try:
            # Use sentence-transformers for encoding
            return self.model.encode(texts, show_progress_bar=True)
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            # Return empty embeddings in case of error
            return np.array([])


class IndexBuilder:
    """Builds and manages FAISS indexes for document retrieval."""
    
    def __init__(self, en_embedder: EnglishEmbedder, zh_embedder: ChineseEmbedder):
        self.en_embedder = en_embedder
        self.zh_embedder = zh_embedder
        
    def build_dual_indexes(self, en_chunks: List[Dict], zh_chunks: List[Dict], index_path: str):
        """Build separate indexes for English and Chinese content."""
        import faiss
        
        # Extract texts and create embeddings
        en_texts = [chunk['text'] for chunk in en_chunks]
        zh_texts = [chunk['text'] for chunk in zh_chunks]
        
        en_embeddings = self.en_embedder.get_embeddings(en_texts) if en_texts else np.array([])
        zh_embeddings = self.zh_embedder.get_embeddings(zh_texts) if zh_texts else np.array([])
        
        # Determine dimensions
        en_dim = en_embeddings.shape[1] if len(en_embeddings) > 0 else 0
        zh_dim = zh_embeddings.shape[1] if len(zh_embeddings) > 0 else 0
        
        # Create indexes
        if len(en_embeddings) > 0:
            en_index = faiss.IndexFlatL2(en_dim)
            en_index.add(en_embeddings.astype('float32'))
            faiss.write_index(en_index, f"{index_path}/en_index.faiss")
            
        if len(zh_embeddings) > 0:
            zh_index = faiss.IndexFlatL2(zh_dim)
            zh_index.add(zh_embeddings.astype('float32'))
            faiss.write_index(zh_index, f"{index_path}/zh_index.faiss")
            
        # Save metadata
        with open(f"{index_path}/en_chunks.json", "w", encoding="utf-8") as f:
            json.dump(en_chunks, f, ensure_ascii=False, indent=2)
            
        with open(f"{index_path}/zh_chunks.json", "w", encoding="utf-8") as f:
            json.dump(zh_chunks, f, ensure_ascii=False, indent=2)
            
        print(f"✅ Built dual language indexes with {len(en_chunks)} English chunks and {len(zh_chunks)} Chinese chunks")


class DocumentProcessor:
    """Main class to process documents and build indexes."""
    
    def __init__(self, index_path: str = "./indexes"):
        self.language_detector = LanguageDetector()
        self.chunker = RecursiveChunker()
        
        # Initialize embedders only when needed to avoid loading models twice
        self._en_embedder = None
        self._zh_embedder = None
        self.index_path = index_path
        
    @property
    def en_embedder(self):
        """Lazy initialization of English embedder"""
        if self._en_embedder is None:
            print("Initializing English embedder...")
            self._en_embedder = EnglishEmbedder()
        return self._en_embedder
    
    @property
    def zh_embedder(self):
        """Lazy initialization of Chinese embedder"""
        if self._zh_embedder is None:
            print("Initializing Chinese embedder...")
            self._zh_embedder = ChineseEmbedder()
        return self._zh_embedder
    
    @property
    def index_builder(self):
        """Lazy initialization of index builder"""
        return IndexBuilder(self.en_embedder, self.zh_embedder)
        
    def process_documents(self, zh_doc_path: str = None, en_doc_path: str = None):
        """Process both language documents and build indexes."""
        import os
        
        # Create index directory if it doesn't exist
        os.makedirs(self.index_path, exist_ok=True)
        
        zh_chunks = []
        en_chunks = []
        
        # Process Chinese document if provided
        if zh_doc_path and os.path.exists(zh_doc_path):
            zh_sections = self.language_detector.process_document(zh_doc_path, 'zh')
            zh_chunks = self.chunker.create_chunks(zh_sections, 'zh')
            print(f"✅ Processed Chinese document: {len(zh_chunks)} chunks created")
        
        # Process English document if provided
        if en_doc_path and os.path.exists(en_doc_path):
            en_sections = self.language_detector.process_document(en_doc_path, 'en')
            en_chunks = self.chunker.create_chunks(en_sections, 'en')
            print(f"✅ Processed English document: {len(en_chunks)} chunks created")
        
        # Build indexes
        self.index_builder.build_dual_indexes(en_chunks, zh_chunks, self.index_path)
        
        return len(zh_chunks), len(en_chunks)
        
    def process_document(self, doc_path: str):
        """Legacy method for backward compatibility."""
        import os
        if not os.path.exists(doc_path):
            print(f"⚠️ Document not found: {doc_path}")
            return 0, 0
            
        # Detect language of the document
        with open(doc_path, 'rb') as f:
            doc = Document(doc_path)
            sample_text = " ".join([p.text for p in doc.paragraphs[:10]])
            language = self.language_detector.detect_language(sample_text)
        
        if language == 'zh':
            return self.process_documents(zh_doc_path=doc_path)
        else:
            return self.process_documents(en_doc_path=doc_path)


# Example usage
if __name__ == "__main__":
    import os
    from dotenv import load_dotenv
    
    # Load environment variables
    load_dotenv()
    
    processor = DocumentProcessor()
    
    # Get document paths from environment variables
    zh_doc_path = os.getenv("HANDBOOK_PATH_ZH")
    en_doc_path = os.getenv("HANDBOOK_PATH_EN")
    
    # Process documents
    zh_count, en_count = processor.process_documents(zh_doc_path, en_doc_path)
    print(f"Processed documents: {zh_count} Chinese chunks, {en_count} English chunks")